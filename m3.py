import streamlit as st
import sqlite3
from PyPDF2 import PdfReader
from docx import Document
import os
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.docstore.document import Document as LangchainDocument
import time
from functools import lru_cache
import google.generativeai as genai
from PIL import Image
import pytube
import speech_recognition as sr

# Cargar variables de entorno y configurar la API de Google
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Función para reconocimiento de voz
def reconocer_voz():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        st.write("Habla ahora...")
        audio = r.listen(source)
        st.write("Procesando...")
    
    try:
        texto = r.recognize_google(audio, language="es-ES")
        return texto
    except sr.UnknownValueError:
        return "No se pudo entender el audio"
    except sr.RequestError as e:
        return f"Error en el servicio de reconocimiento de voz: {e}"
    
# Nueva función para procesar videos de YouTube
def procesar_video_youtube(url_video):
    try:
        yt = pytube.YouTube(url_video)
        video = yt.streams.filter(progressive=True, file_extension='mp4').order_by('resolution').desc().first()
        
        if not video:
            return "No se pudo encontrar un stream de video adecuado."
        
        ruta_video = video.download(output_path="temp")
        
        modelo = genai.GenerativeModel('gemini-1.5-flash')
        respuesta = modelo.generate_content([
            "Analiza este video relacionado con ingeniería de sonido o equipos de audio. Describe lo que ves y escuchas en detalle, enfocándote en aspectos técnicos que puedan ser relevantes para una consulta de ingeniería de sonido.",
            genai.types.ImageInput(ruta_video)
        ])
        
        os.remove(ruta_video)
        return respuesta.text
    except Exception as e:
        return f"Error al procesar el video: {str(e)}"
    
# Función para obtener embeddings (en caché)
@st.cache_resource
def obtener_embeddings():
    return GoogleGenerativeAIEmbeddings(model="models/embedding-001")

# Conexión a la base de datos en caché
@lru_cache(maxsize=None)
def obtener_conexion_bd():
    conn = sqlite3.connect('base_datos_ingenieria_sonido.db')
    conn.execute('''CREATE TABLE IF NOT EXISTS documentos
                    (id INTEGER PRIMARY KEY, nombre TEXT UNIQUE, contenido TEXT)''')
    return conn

# Funciones de la base de datos (sin cambios)
def documento_existe(conn, nombre):
    return conn.execute("SELECT 1 FROM documentos WHERE nombre = ?", (nombre,)).fetchone() is not None

def guardar_documento(conn, nombre, contenido):
    conn.execute("INSERT OR REPLACE INTO documentos (nombre, contenido) VALUES (?, ?)", (nombre, contenido))
    conn.commit()

def obtener_todos_documentos(conn):
    return conn.execute("SELECT nombre, contenido FROM documentos ORDER BY nombre ASC").fetchall()

def eliminar_documento(conn, nombre):
    conn.execute("DELETE FROM documentos WHERE nombre = ?", (nombre,))
    conn.commit()

# Funciones para procesar archivos (sin cambios)
@st.cache_data
def obtener_texto_archivo(archivo):
    if archivo.name.lower().endswith('.pdf'):
        return obtener_texto_pdf(archivo)
    elif archivo.name.lower().endswith('.docx'):
        return obtener_texto_docx(archivo)
    elif archivo.name.lower().endswith('.txt'):
        return archivo.getvalue().decode("utf-8")
    return ""

@st.cache_data
def obtener_texto_pdf(archivo_pdf):
    texto = ""
    pdf_reader = PdfReader(archivo_pdf)
    for pagina in pdf_reader.pages:
        texto += pagina.extract_text()
    return texto

@st.cache_data
def obtener_texto_docx(archivo_docx):
    doc = Document(archivo_docx)
    return "\n".join(parrafo.text for parrafo in doc.paragraphs)

# Función para procesar imagen y extraer texto (sin cambios)
def procesar_imagen_a_texto(archivo_imagen):
    imagen = Image.open(archivo_imagen)
    modelo = genai.GenerativeModel('gemini-1.5-flash')
    respuesta = modelo.generate_content([
        "Analiza esta imagen relacionada con ingeniería de sonido o equipos de audio. Describe lo que ves en detalle, enfocándote en aspectos técnicos que puedan ser relevantes para una consulta de ingeniería de sonido.",
        imagen
    ])
    return respuesta.text

# Función para verificar la clave de acceso (sin cambios)
def verificar_clave():
    clave = st.text_input("Ingrese la clave de acceso:", type="password")
    return clave == "matusay"  # Reemplazar con un método seguro en producción

# Nueva función para el repensamiento (Chain of Thought)
# Función modificada para el repensamiento (Chain of Thought)
# Función modificada para el repensamiento (Cadena de Pensamiento)
def repensar_respuesta(respuesta_inicial, consulta, descripcion_imagen, descripcion_video, contexto):
    plantilla_repensamiento = """
    Como experto en ingeniería de sonido, analiza críticamente la siguiente respuesta inicial a una consulta. 
    Mejora la respuesta considerando el contexto adicional, la descripción de la imagen y la descripción del video si están disponibles.
    No expliques el proceso de mejora, simplemente proporciona la respuesta mejorada.

    Consulta original: {consulta}
    Descripción de la imagen (si está disponible): {descripcion_imagen}
    Descripción del video (si está disponible): {descripcion_video}
    Contexto adicional: {context}
    
    Respuesta inicial: {respuesta_inicial}
    
    Proporciona una respuesta mejorada, más completa y específica para la consulta de ingeniería de sonido:
    """
    
    modelo = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.2)
    prompt = PromptTemplate(template=plantilla_repensamiento, input_variables=["consulta", "descripcion_imagen", "descripcion_video", "context", "respuesta_inicial"])
    
    cadena_repensamiento = load_qa_chain(modelo, chain_type="stuff", prompt=prompt)
    
    respuesta_mejorada = cadena_repensamiento.invoke({
        "input_documents": [LangchainDocument(page_content=contexto)],
        "consulta": consulta,
        "descripcion_imagen": descripcion_imagen,
        "descripcion_video": descripcion_video,
        "context": contexto,
        "respuesta_inicial": respuesta_inicial
    })["output_text"]
    
    return respuesta_mejorada

# Función modificada para obtener la cadena del asistente de ingeniería de sonido
@lru_cache(maxsize=1)
def obtener_cadena_ingenieria_sonido():
    plantilla_prompt = """
    Eres un asistente experto en ingeniería de sonido. Utiliza tu conocimiento para responder la siguiente pregunta.
    Si se proporciona una imagen, utiliza la descripción de la imagen para dar contexto a tu respuesta.
    
    Pregunta: {question}
    
    Descripción de la imagen (si está disponible): {image_description}
    
    Información relevante de la base de conocimientos:
    {context}
    
    Proporciona una respuesta detallada y técnicamente precisa:
    """
    modelo = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.4)
    prompt = PromptTemplate(template=plantilla_prompt, input_variables=["question", "image_description", "context"])
    return load_qa_chain(modelo, chain_type="stuff", prompt=prompt)

# Función modificada para responder consultas de ingeniería de sonido
def responder_consulta(consulta, descripcion_imagen, descripcion_video, documentos, historial_conversacion):
    cadena = obtener_cadena_ingenieria_sonido()
    contexto = "\n".join([doc[1] for doc in documentos])
    doc = LangchainDocument(page_content=contexto)
    
    historial_formateado = "\n".join([f"Humano: {q}\nAsistente: {a}" for q, a in historial_conversacion])
    
    plantilla_prompt = """
    Eres un asistente experto en ingeniería de sonido. Utiliza tu conocimiento para responder la siguiente pregunta.
    Si se proporciona una imagen o un video, utiliza sus descripciones para dar contexto a tu respuesta.
    
    Historial de la conversación:
    {historial}
    
    Pregunta actual: {question}
    
    Descripción de la imagen (si está disponible): {image_description}
    
    Descripción del video (si está disponible): {video_description}
    
    Información relevante de la base de conocimientos:
    {context}
    
    Proporciona una respuesta detallada y técnicamente precisa, teniendo en cuenta el contexto de la conversación previa, la imagen y el video si están disponibles:
    """
    
    prompt = PromptTemplate(template=plantilla_prompt, input_variables=["historial", "question", "image_description", "video_description", "context"])
    
    cadena_actualizada = load_qa_chain(ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.4), chain_type="stuff", prompt=prompt)
    
    respuesta_inicial = cadena_actualizada.invoke({
        "input_documents": [doc],
        "historial": historial_formateado,
        "question": consulta,
        "image_description": descripcion_imagen,
        "video_description": descripcion_video
    })["output_text"]
    
    respuesta_final = repensar_respuesta(respuesta_inicial, consulta, descripcion_imagen, descripcion_video, contexto)
    
    return respuesta_final, respuesta_inicial

# Función principal de la aplicación (sin cambios mayores)
def main():
    st.set_page_config(page_title="Asistente de Ingeniería de Sonido", layout="wide")

    st.markdown(
        """
        <h1 style="text-align: center; animation: colorChange 5s infinite;">
        Bienvenido al Asistente de Ingeniería de Sonido
        </h1>
        <h3 style='text-align: center; color: #888888;'>Sube documentos y haz preguntas técnicas sobre ingeniería de sonido</h3>
        """,
        unsafe_allow_html=True
    )

    conn = obtener_conexion_bd()

    # Inicializar el estado de la sesión si no existe
    if 'historial_conversacion' not in st.session_state:
        st.session_state.historial_conversacion = []
    if 'respuesta_final' not in st.session_state:
        st.session_state.respuesta_final = ""
    if 'respuesta_inicial' not in st.session_state:
        st.session_state.respuesta_inicial = ""

    menu = ["Inicio", "Ver Documentos", "Subir Documentos", "Hacer Preguntas", "Configuración"]
    eleccion = st.sidebar.selectbox("Menú", menu)

    if eleccion == "Inicio":
        st.write("Bienvenido al Asistente de Ingeniería de Sonido. Usa el menú de la izquierda para navegar.")

    elif eleccion == "Ver Documentos":
        st.subheader("Documentos Subidos")
        documentos = obtener_todos_documentos(conn)
        if documentos:
            for doc in documentos:
                st.write(f"- {doc[0]}")
        else:
            st.info("No hay documentos en la base de datos. Por favor, sube algunos documentos primero.")

    elif eleccion == "Subir Documentos":
        st.subheader("Subir Nuevos Documentos")
        archivos_subidos = st.file_uploader("Sube archivos PDF, DOCX o TXT", accept_multiple_files=True, type=['pdf', 'docx', 'txt'])
        
        if st.button("Procesar Archivos"):
            if archivos_subidos:
                with st.spinner("Procesando archivos... Por favor espera."):
                    for archivo in archivos_subidos:
                        if not documento_existe(conn, archivo.name):
                            contenido = obtener_texto_archivo(archivo)
                            guardar_documento(conn, archivo.name, contenido)
                            st.success(f"Documento '{archivo.name}' procesado y guardado.")
                        else:
                            st.warning(f"El documento '{archivo.name}' ya existe en la base de datos.")
                st.success("Todos los documentos han sido procesados.")
            else:
                st.warning("Por favor, sube al menos un archivo antes de procesar.")

    elif eleccion == "Hacer Preguntas":
        st.subheader("Hacer Preguntas de Ingeniería de Sonido")
        
        # Mostrar el historial de la conversación
        if st.session_state.historial_conversacion:
            st.write("Historial de la conversación:")
            for q, a in st.session_state.historial_conversacion:
                st.text_area("Pregunta:", value=q, height=100, disabled=True)
                st.text_area("Respuesta:", value=a, height=200, disabled=True)
                st.markdown("---")
        else:
            st.info("No hay historial de conversación. Haz una pregunta para comenzar.")
        
        # Opciones para ingresar la consulta
        metodo_entrada = st.radio("Elige cómo quieres hacer tu pregunta:", ("Texto", "Voz"))
        
        if metodo_entrada == "Texto":
            consulta = st.text_area("Ingresa tu pregunta sobre ingeniería de sonido:")
        else:
            if st.button("Iniciar grabación de voz"):
                consulta = reconocer_voz()
                st.write(f"Pregunta reconocida: {consulta}")
        
        imagen_subida = st.file_uploader("Sube una imagen (opcional)", type=['png', 'jpg', 'jpeg'])
        url_youtube = st.text_input("Ingresa una URL de YouTube (opcional):")
        
        if st.button("Obtener Respuesta"):
            if consulta:
                with st.spinner("Analizando y generando respuesta detallada..."):
                    descripcion_imagen = ""
                    if imagen_subida:
                        descripcion_imagen = procesar_imagen_a_texto(imagen_subida)
                    
                    descripcion_video = ""
                    if url_youtube:
                        descripcion_video = procesar_video_youtube(url_youtube)
                    
                    documentos = obtener_todos_documentos(conn)
                    if documentos:
                        st.session_state.respuesta_final, st.session_state.respuesta_inicial = responder_consulta(consulta, descripcion_imagen, descripcion_video, documentos, st.session_state.historial_conversacion)
                        st.write("Respuesta:", st.session_state.respuesta_final)
                        
                        # Agregar la nueva pregunta y respuesta al historial
                        st.session_state.historial_conversacion.append((consulta, st.session_state.respuesta_final))
                    else:
                        st.warning("No hay documentos en la base de datos. Por favor, sube algunos documentos primero.")
            else:
                st.warning("Por favor, ingresa una pregunta antes de enviar.")

        # Botón para mostrar el proceso de pensamiento
        if st.button("Mostrar proceso de pensamiento"):
            if st.session_state.respuesta_inicial and st.session_state.respuesta_final:
                st.write("Respuesta inicial:", st.session_state.respuesta_inicial)
                st.write("Respuesta mejorada:", st.session_state.respuesta_final)
            else:
                st.info("Aún no hay un proceso de pensamiento para mostrar. Por favor, haz una pregunta primero.")

        # Opción para limpiar el historial
        if st.button("Limpiar Historial"):
            st.session_state.historial_conversacion = []
            st.session_state.respuesta_final = ""
            st.session_state.respuesta_inicial = ""
            st.success("Historial de conversación limpiado.")
            st.rerun()  # Forzar la recarga de la página

    elif eleccion == "Configuración":
        if verificar_clave():
            st.subheader("Configuración y Respaldo de Base de Datos")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Descargar Base de Datos")
                if st.button("Descargar BD"):
                    with open('base_datos_ingenieria_sonido.db', 'rb') as f:
                        datos_bytes = f.read()
                    st.download_button(
                        label="Descargar base_datos_ingenieria_sonido.db",
                        data=datos_bytes,
                        file_name="base_datos_ingenieria_sonido.db",
                        mime="application/octet-stream"
                    )
            
            with col2:
                st.subheader("Subir Base de Datos")
                archivo_subido = st.file_uploader("Selecciona el archivo de base de datos", type=['db'])
                if archivo_subido is not None:
                    if st.button("Subir BD"):
                        with open('base_datos_ingenieria_sonido.db', 'wb') as f:
                            f.write(archivo_subido.getbuffer())
                        st.success("Base de datos subida exitosamente. Reinicia la aplicación para ver los cambios.")
        else:
            st.error("Clave de acceso incorrecta. Acceso denegado.")

if __name__ == "__main__":
    main()
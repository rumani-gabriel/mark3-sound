import streamlit as st
import sqlite3
from PyPDF2 import PdfReader
from docx import Document
import os
import json
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from langchain.docstore.document import Document as LangchainDocument
import google.generativeai as genai
from PIL import Image
import speech_recognition as sr

# Configuración de la página Streamlit
st.set_page_config(page_title="Chatbot con Base de Conocimientos", layout="wide")

# Cargar variables de entorno y configurar la API de Google
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Configurar rutas para almacenamiento persistente
DB_PATH = "knowledge_base.db"
VECTORSTORE_PATH = "vectorstore.faiss"

# Inicializar la base de datos SQLite persistente
@st.cache_resource
def get_database_connection():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.execute('''CREATE TABLE IF NOT EXISTS documents
                    (id INTEGER PRIMARY KEY, name TEXT UNIQUE, content TEXT)''')
    conn.commit()
    return conn

conn = get_database_connection()

# Inicializar o cargar el vector store
@st.cache_resource
def load_vectorstore():
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    if os.path.exists(VECTORSTORE_PATH):
        try:
            return FAISS.load_local(VECTORSTORE_PATH, embeddings, allow_dangerous_deserialization=True)
        except Exception as e:
            st.error(f"Error al cargar el vectorstore: {str(e)}")
            st.warning("Creando un nuevo vectorstore vacío.")
            return FAISS.from_texts([""], embeddings)
    return FAISS.from_texts([""], embeddings)

vectorstore = load_vectorstore()

# Funciones para procesar archivos
@st.cache_data
def get_file_text(file):
    if file.name.lower().endswith('.pdf'):
        return get_pdf_text(file)
    elif file.name.lower().endswith('.docx'):
        return get_docx_text(file)
    return file.getvalue().decode()

@st.cache_data
def get_pdf_text(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

@st.cache_data
def get_docx_text(docx_file):
    doc = Document(docx_file)
    return "\n".join(para.text for para in doc.paragraphs)

# Función para procesar imagen y extraer texto
def process_image_to_text(image_file):
    image = Image.open(image_file)
    model = genai.GenerativeModel('gemini-pro-vision')
    response = model.generate_content(["Describe detalladamente lo que ves en esta imagen.", image])
    return response.text

# Función para procesar audio usando reconocimiento de voz
def process_audio_to_text():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        st.write("Habla ahora...")
        audio = r.listen(source)
        try:
            text = r.recognize_google(audio, language="es-ES")
            st.session_state.user_input = text
            st.write("Has dicho: " + text)
            return text
        except sr.UnknownValueError:
            st.write("No entendí lo que dijiste.")
            return ""
        except sr.RequestError as e:
            st.write("Hubo un error al procesar tu audio: {0}".format(e))
            return ""

# Funciones para la base de conocimientos
def add_to_knowledge_base(conn, name, content):
    global vectorstore
    
    try:
        # Agregar el documento a la base de datos SQLite
        conn.execute("INSERT OR REPLACE INTO documents (name, content) VALUES (?, ?)", (name, content))
        conn.commit()

        # Actualizar el vector store
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        
        # Dividir el contenido en chunks más pequeños
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = text_splitter.split_text(content)
        
        # Agregar los nuevos chunks al vectorstore existente
        vectorstore.add_texts(chunks)
        
        # Guardar el vectorstore actualizado
        vectorstore.save_local(VECTORSTORE_PATH)
        
        st.success(f"Documento '{name}' agregado exitosamente a la base de conocimientos.")
    except Exception as e:
        st.error(f"Error al agregar documento: {str(e)}")

def get_all_documents(conn):
    return conn.execute("SELECT name, content FROM documents").fetchall()

# Configuración de Langchain y el modelo
@st.cache_resource
def setup_qa_chain():
    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)
    prompt_template = """
    Utiliza la siguiente información para responder a la pregunta del usuario.
    Si no conoces la respuesta, simplemente di que no lo sabes.

    Contexto: {context}

    Pregunta del usuario: {question}

    Respuesta:
    """
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    return load_qa_chain(model, chain_type="stuff", prompt=prompt)

# Función para generar respuesta
def generate_response(user_input, image_context=""):
    try:
        if vectorstore:
            relevant_docs = vectorstore.similarity_search(user_input, k=3)
            context = "\n".join([doc.page_content for doc in relevant_docs])
        else:
            context = ""

        combined_context = context + "\n" + image_context

        chain = setup_qa_chain()
        response = chain.invoke({
            "input_documents": [LangchainDocument(page_content=combined_context)],
            "question": user_input
        })

        return response["output_text"]
    except Exception as e:
        st.error(f"Error al generar respuesta: {str(e)}")
        return "Lo siento, hubo un error al procesar tu pregunta."

# Función principal de la aplicación
def main():
    st.title("Chatbot con Base de Conocimientos")

    # Sidebar para cargar documentos
    st.sidebar.header("Cargar Documentos")
    uploaded_file = st.sidebar.file_uploader("Sube un archivo (PDF, DOCX, TXT)", type=['pdf', 'docx', 'txt'])
    if uploaded_file:
        with st.spinner("Procesando documento..."):
            content = get_file_text(uploaded_file)
            add_to_knowledge_base(conn, uploaded_file.name, content)

    # Mostrar documentos cargados
    st.sidebar.header("Documentos Cargados")
    documents = get_all_documents(conn)
    for doc in documents:
        st.sidebar.text(doc[0])

    # Inicializar el historial de chat y user_input en session_state
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'user_input' not in st.session_state:
        st.session_state.user_input = ""

    # Área principal de chat
    st.header("Chat")
    for message in st.session_state.chat_history:
        st.write(f"{'Usuario' if message['is_user'] else 'Asistente'}: {message['text']}")

    # Input del usuario
    user_input = st.text_input("Escribe tu pregunta aquí:", value=st.session_state.user_input, key="text_input")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("🎙️ Hablar"):
            process_audio_to_text()
            st.rerun()
    
    with col2:
        uploaded_image = st.file_uploader("Sube una imagen para dar contexto (opcional)", type=['png', 'jpg', 'jpeg'])

    with col3:
        if st.button("Enviar"):
            if user_input:
                st.session_state.chat_history.append({"text": user_input, "is_user": True})
                
                image_context = ""
                if uploaded_image:
                    image_context = process_image_to_text(uploaded_image)

                with st.spinner("Generando respuesta..."):
                    response = generate_response(user_input, image_context)

                st.session_state.chat_history.append({"text": response, "is_user": False})
                st.write(f"Asistente: {response}")
                
                st.session_state.user_input = ""
                st.rerun()

if __name__ == "__main__":
    main()